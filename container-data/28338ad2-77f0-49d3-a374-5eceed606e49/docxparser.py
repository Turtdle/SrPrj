import docx
import argparse
from pathlib import Path
import builtins

def extract_text_from_paragraph(paragraph):
    text = paragraph.text
    if text.strip():  
        if paragraph.style.name.startswith('Heading'):
            return f"[{paragraph.style.name}] {text}\n"
        return text + "\n"
    return "\n"  

def extract_text_from_table(table):
    table_text = "[TABLE]\n"
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell_text = " ".join([p.text for p in cell.paragraphs])
            table_text += f"[Cell {i},{j}] {cell_text}\n"
    table_text += "[/TABLE]\n"
    return table_text

def parse_docx(docx_path, output_path=None):
    try:
        doc = docx.Document(docx_path)
        content = []
        
        # Document properties
        properties = doc.core_properties
        content.append(f"Document Properties:")
        content.append(f"Title: {properties.title or 'Not specified'}")
        content.append(f"Author: {properties.author or 'Not specified'}")
        content.append(f"Created: {properties.created or 'Not specified'}")
        content.append(f"Modified: {properties.modified or 'Not specified'}")
        content.append("\n" + "="*50 + "\n")
        
    
        for element in doc.element.body:
            if element.tag.endswith('p'):
                for paragraph in doc.paragraphs:
                    if paragraph._element is element:
                        content.append(extract_text_from_paragraph(paragraph))
                        break
            elif element.tag.endswith('tbl'): 
                for table in doc.tables:
                    if table._element is element:
                        content.append(extract_text_from_table(table))
                        break
        
        result = "".join(content)
        
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result)
            print(f"Parsed content saved to {output_path}")
        
        return result
        
    except Exception as e:
        print(f"Error parsing DOCX file: {e}")
        return None

